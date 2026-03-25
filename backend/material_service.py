"""
EduFlow Material Analysis Service
Handles file upload, parsing, and AI-powered analysis of teaching materials
"""

import os
import json
import logging
import base64
from typing import Optional, Dict, Any
from io import BytesIO
from pathlib import Path
from dotenv import load_dotenv

# Load environment variables
ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

import openai
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from PIL import Image

logger = logging.getLogger(__name__)

# Initialize OpenAI client (after env is loaded)
openai_client = openai.OpenAI(api_key=os.environ.get('OPENAI_API_KEY'))

# Supported file types
SUPPORTED_EXTENSIONS = {
    'pdf': 'application/pdf',
    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'doc': 'application/msword',
    'txt': 'text/plain',
    'png': 'image/png',
    'jpg': 'image/jpeg',
    'jpeg': 'image/jpeg',
}

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB


class MaterialParser:
    """Parse different file types and extract text content"""
    
    @staticmethod
    def parse_pdf(file_bytes: bytes) -> Dict[str, Any]:
        """Extract text from PDF"""
        try:
            reader = PdfReader(BytesIO(file_bytes))
            pages = []
            full_text = ""
            
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                pages.append({
                    "page_number": i + 1,
                    "text": text
                })
                full_text += text + "\n\n"
            
            return {
                "success": True,
                "file_type": "pdf",
                "page_count": len(reader.pages),
                "pages": pages,
                "full_text": full_text.strip(),
                "char_count": len(full_text),
                "word_count": len(full_text.split())
            }
        except Exception as e:
            logger.error(f"PDF parsing error: {e}")
            return {"success": False, "error": str(e)}
    
    @staticmethod
    def parse_docx(file_bytes: bytes) -> Dict[str, Any]:
        """Extract text from Word document"""
        try:
            doc = DocxDocument(BytesIO(file_bytes))
            paragraphs = []
            full_text = ""
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
                    full_text += para.text + "\n"
            
            # Also extract tables
            tables_text = ""
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    tables_text += row_text + "\n"
            
            full_text += tables_text
            
            return {
                "success": True,
                "file_type": "docx",
                "paragraph_count": len(paragraphs),
                "paragraphs": paragraphs[:50],  # First 50 paragraphs
                "full_text": full_text.strip(),
                "char_count": len(full_text),
                "word_count": len(full_text.split()),
                "has_tables": len(doc.tables) > 0
            }
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            return {"success": False, "error": str(e)}
    
    @staticmethod
    def parse_txt(file_bytes: bytes) -> Dict[str, Any]:
        """Extract text from plain text file"""
        try:
            text = file_bytes.decode('utf-8', errors='ignore')
            lines = text.split('\n')
            
            return {
                "success": True,
                "file_type": "txt",
                "line_count": len(lines),
                "full_text": text.strip(),
                "char_count": len(text),
                "word_count": len(text.split())
            }
        except Exception as e:
            logger.error(f"TXT parsing error: {e}")
            return {"success": False, "error": str(e)}
    
    @staticmethod
    def parse_image(file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Process image file"""
        try:
            img = Image.open(BytesIO(file_bytes))
            
            # Convert to base64 for vision API
            buffered = BytesIO()
            img_format = img.format or 'PNG'
            img.save(buffered, format=img_format)
            img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
            
            return {
                "success": True,
                "file_type": "image",
                "format": img_format.lower(),
                "width": img.width,
                "height": img.height,
                "mode": img.mode,
                "base64": img_base64,
                "full_text": "[Bild - Text wird mit Vision-API extrahiert]"
            }
        except Exception as e:
            logger.error(f"Image parsing error: {e}")
            return {"success": False, "error": str(e)}
    
    @classmethod
    def parse(cls, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Parse file based on extension"""
        ext = filename.lower().split('.')[-1] if '.' in filename else ''
        
        if ext == 'pdf':
            return cls.parse_pdf(file_bytes)
        elif ext in ['docx', 'doc']:
            return cls.parse_docx(file_bytes)
        elif ext == 'txt':
            return cls.parse_txt(file_bytes)
        elif ext in ['png', 'jpg', 'jpeg']:
            return cls.parse_image(file_bytes, filename)
        else:
            return {"success": False, "error": f"Nicht unterstützter Dateityp: {ext}"}


class MaterialAnalyzer:
    """AI-powered analysis of teaching materials"""
    
    @staticmethod
    async def analyze_text_content(text: str, filename: str) -> Dict[str, Any]:
        """Analyze text content to extract educational metadata"""
        
        # Truncate very long texts
        max_chars = 8000
        truncated = text[:max_chars] if len(text) > max_chars else text
        
        system_prompt = """Du bist ein Experte für Schweizer Schulunterricht und Lehrplananalyse. Analysiere Unterrichtsmaterialien präzise.

Analysiere das folgende Unterrichtsmaterial und extrahiere die Informationen im JSON-Format:
{
  "detected_subject": "Fach (Deutsch/Mathematik/NMG/Englisch/Französisch/Andere)",
  "detected_topic": "Hauptthema des Materials",
  "detected_grade": "Geschätzte Klassenstufe (1-6 oder 'unklar')",
  "document_type": "Art des Dokuments (Arbeitsblatt/Prüfung/Lehrbuchseite/Notizen/Übung/Andere)",
  "keywords": ["Schlüsselwort1", "Schlüsselwort2", "..."],
  "learning_goals": ["Mögliches Lernziel 1", "Mögliches Lernziel 2"],
  "difficulty_estimate": "easy/medium/hard",
  "content_summary": "Kurze Zusammenfassung des Inhalts (2-3 Sätze)",
  "suggested_actions": ["Arbeitsblatt erstellen", "Quiz erstellen", "Vereinfachen"],
  "language_level": "Sprachniveau (einfach/mittel/anspruchsvoll)",
  "has_exercises": true/false,
  "has_solutions": true/false,
  "source_quality": "Qualitätseinschätzung (gut/mittel/überarbeiten)"
}

Antworte NUR mit dem JSON-Objekt, ohne zusätzlichen Text."""

        try:
            response = openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"Dateiname: {filename}\n\nInhalt:\n{truncated}"}
                ],
                temperature=0.3,
                max_tokens=1000,
                response_format={"type": "json_object"}
            )
            
            result = json.loads(response.choices[0].message.content)
            result["analysis_success"] = True
            return result
            
        except Exception as e:
            logger.error(f"Material analysis error: {e}")
            return {
                "analysis_success": False,
                "error": str(e),
                "detected_subject": "Unbekannt",
                "detected_topic": filename,
                "detected_grade": "unklar",
                "document_type": "Unbekannt",
                "keywords": [],
                "learning_goals": [],
                "content_summary": "Analyse fehlgeschlagen",
                "suggested_actions": ["Manuell prüfen"]
            }
    
    @staticmethod
    async def analyze_image_content(base64_image: str, filename: str) -> Dict[str, Any]:
        """Analyze image content using vision API"""
        
        try:
            response = openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": "Du bist ein Experte für Schweizer Schulunterricht. Analysiere Bilder von Unterrichtsmaterialien."
                    },
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": """Analysiere dieses Bild eines Unterrichtsmaterials und gib folgende Informationen als JSON zurück:
{
  "detected_subject": "Fach",
  "detected_topic": "Thema",
  "detected_grade": "Klassenstufe (1-6 oder unklar)",
  "document_type": "Art des Dokuments",
  "extracted_text": "Wichtiger Text aus dem Bild",
  "keywords": ["Schlüsselwörter"],
  "content_summary": "Beschreibung des Bildinhalts",
  "suggested_actions": ["Mögliche Aktionen"],
  "is_worksheet": true/false,
  "has_handwriting": true/false
}"""
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{base64_image}",
                                    "detail": "high"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=1000,
                response_format={"type": "json_object"}
            )
            
            result = json.loads(response.choices[0].message.content)
            result["analysis_success"] = True
            result["is_image"] = True
            return result
            
        except Exception as e:
            logger.error(f"Image analysis error: {e}")
            return {
                "analysis_success": False,
                "error": str(e),
                "is_image": True,
                "detected_subject": "Unbekannt",
                "content_summary": "Bildanalyse fehlgeschlagen"
            }


class MaterialTransformer:
    """Transform uploaded materials into new educational content"""
    
    @staticmethod
    def get_transformation_prompt(action: str, source_text: str, options: Dict[str, Any]) -> str:
        """Generate prompt for material transformation"""
        
        grade = options.get('grade', '4')
        subject = options.get('subject', 'Deutsch')
        difficulty = options.get('difficulty', 'medium')
        question_count = options.get('questionCount', 10)
        custom_instructions = options.get('customInstructions', '')
        
        difficulty_map = {
            'easy': 'Einfach - grundlegende Konzepte, einfache Sprache',
            'medium': 'Mittel - altersgerechte Herausforderung', 
            'hard': 'Schwierig - anspruchsvolle Aufgaben'
        }
        
        base_context = f"""
Du erstellst Unterrichtsmaterial für die Schweizer Primarschule.
Klassenstufe: {grade}. Klasse
Fach: {subject}
Schwierigkeit: {difficulty_map.get(difficulty, 'Mittel')}
"""

        # Action-specific prompts
        action_prompts = {
            'worksheet': f"""Erstelle ein Arbeitsblatt mit {question_count} Aufgaben basierend auf dem folgenden Material.
Das Arbeitsblatt soll:
- Abwechslungsreiche Aufgabentypen enthalten
- Dem Niveau der {grade}. Klasse entsprechen
- Klare Anweisungen haben
- Lösungen für die Lehrperson enthalten

{custom_instructions}

Quellmaterial:
{source_text}""",
            
            'exam': f"""Erstelle eine Prüfung mit {question_count} Aufgaben basierend auf dem folgenden Material.
Die Prüfung soll:
- Verschiedene Schwierigkeitsgrade abdecken
- Punkte pro Aufgabe angeben (Total: ca. {question_count * 3} Punkte)
- Lösungsschlüssel enthalten
- Bewertungskriterien haben

{custom_instructions}

Quellmaterial:
{source_text}""",

            'summary': f"""Erstelle eine schülergerechte Zusammenfassung des folgenden Materials für die {grade}. Klasse.
Die Zusammenfassung soll:
- Die wichtigsten Punkte hervorheben
- Einfache, klare Sprache verwenden
- In Abschnitte gegliedert sein

{custom_instructions}

Quellmaterial:
{source_text}""",

            'cloze': f"""Erstelle einen Lückentext basierend auf dem folgenden Material.
Der Lückentext soll:
- Ca. 8-12 sinnvolle Lücken haben
- Schlüsselbegriffe als Lücken verwenden
- Eine Wortbank mit den Lösungen enthalten

{custom_instructions}

Quellmaterial:
{source_text}""",

            'quiz': f"""Erstelle ein Quiz mit {question_count} Fragen basierend auf dem folgenden Material.
Das Quiz soll:
- Multiple-Choice und Richtig/Falsch Fragen enthalten
- Schnell zu beantworten sein
- Die Hauptpunkte abfragen

{custom_instructions}

Quellmaterial:
{source_text}""",

            'simplify': f"""Vereinfache das folgende Material für schwächere Schüler der {grade}. Klasse.
Dabei soll:
- Einfachere Wörter verwendet werden
- Kürzere Sätze formuliert werden
- Der Inhalt erhalten bleiben
- Zusätzliche Erklärungen hinzugefügt werden

{custom_instructions}

Quellmaterial:
{source_text}""",

            'differentiate': f"""Erstelle drei differenzierte Versionen des folgenden Materials:
1. EINFACH (Basisniveau)
2. STANDARD (Regelniveau)  
3. ERWEITERT (Erweiterungsniveau)

Jede Version soll für die {grade}. Klasse geeignet sein.

{custom_instructions}

Quellmaterial:
{source_text}""",

            'solutions': f"""Erstelle vollständige Lösungen und Lehrernotizen für das folgende Material.
Enthalten soll:
- Musterlösungen zu allen Aufgaben
- Erklärungen für häufige Schülerfehler
- Tipps für die Unterrichtsdurchführung
- Bewertungshinweise

{custom_instructions}

Quellmaterial:
{source_text}""",

            'reading_comprehension': f"""Erstelle Leseverständnisaufgaben basierend auf dem folgenden Text.
Erstelle:
- Verständnisfragen zum Text
- Interpretationsfragen
- Transferaufgaben
- Wortschatzübungen

{custom_instructions}

Quellmaterial:
{source_text}""",

            'vocabulary': f"""Erstelle Wortschatzübungen basierend auf dem folgenden Material.
Enthalten soll:
- Definitionen wichtiger Begriffe
- Lückentexte mit Vokabeln
- Zuordnungsübungen
- Satzbildungsübungen

{custom_instructions}

Quellmaterial:
{source_text}"""
        }
        
        return base_context + "\n\n" + action_prompts.get(action, action_prompts['worksheet'])
    
    @staticmethod
    async def transform(
        source_text: str,
        action: str,
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Transform source material into new content"""
        
        prompt = MaterialTransformer.get_transformation_prompt(action, source_text, options)
        
        # JSON output format
        output_format = """
Antworte im folgenden JSON-Format:
{
  "title": "Titel des erstellten Materials",
  "questions": [
    {
      "id": "q1",
      "number": 1,
      "type": "multiple_choice|short_answer|gap_text|true_false|long_answer",
      "question": "Fragetext",
      "options": ["Option A", "Option B", "..."] (nur für Multiple Choice),
      "answer": "Korrekte Antwort",
      "explanation": "Erklärung für Lehrperson",
      "points": 2 (für Prüfungen),
      "answerLines": 3,
      "source_reference": "Basiert auf Abschnitt X des Quellmaterials"
    }
  ],
  "teacher_notes": "Hinweise für die Lehrperson",
  "total_points": 30 (für Prüfungen),
  "estimated_time": "20-30 Minuten",
  "source_summary": "Kurze Beschreibung des verwendeten Quellmaterials"
}"""

        try:
            response = openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "Du bist ein erfahrener Schweizer Primarschullehrer, der hochwertige Unterrichtsmaterialien erstellt."},
                    {"role": "user", "content": prompt + "\n\n" + output_format}
                ],
                temperature=0.7,
                max_tokens=4000,
                response_format={"type": "json_object"}
            )
            
            result = json.loads(response.choices[0].message.content)
            result["transform_success"] = True
            result["action"] = action
            result["source_based"] = True
            
            # Ensure question IDs
            for i, q in enumerate(result.get("questions", [])):
                if "id" not in q:
                    q["id"] = f"q{i+1}"
                if "number" not in q:
                    q["number"] = i + 1
                if "answerLines" not in q:
                    q["answerLines"] = 3
                q["from_source"] = True
            
            return result
            
        except Exception as e:
            logger.error(f"Transformation error: {e}")
            return {
                "transform_success": False,
                "error": str(e)
            }


class ContextualEditor:
    """AI-assisted contextual editing"""
    
    @staticmethod
    async def edit_content(
        current_content: str,
        instruction: str,
        context: Dict[str, Any],
        selection: Optional[str] = None
    ) -> Dict[str, Any]:
        """Apply AI edit to content based on instruction"""
        
        source_info = ""
        if context.get("source_material"):
            source_info = f"\n\nQuellmaterial zur Referenz:\n{context['source_material'][:2000]}"
        
        if selection:
            edit_prompt = f"""Bearbeite den folgenden ausgewählten Text gemäss der Anweisung.

Ausgewählter Text:
"{selection}"

Anweisung: {instruction}

Kontext des gesamten Dokuments:
{current_content[:1500]}
{source_info}

Antworte im JSON-Format:
{{
  "edited_text": "Der bearbeitete Text",
  "explanation": "Kurze Erklärung der Änderungen",
  "changes_made": ["Änderung 1", "Änderung 2"]
}}"""
        else:
            edit_prompt = f"""Bearbeite das folgende Dokument gemäss der Anweisung.

Dokument:
{current_content}

Anweisung: {instruction}
{source_info}

Antworte im JSON-Format:
{{
  "edited_content": "Das vollständige bearbeitete Dokument",
  "explanation": "Kurze Erklärung der Änderungen",
  "changes_made": ["Änderung 1", "Änderung 2"]
}}"""

        try:
            response = openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "Du bist ein Experte für Unterrichtsmaterialien und hilfst beim Bearbeiten und Verbessern von Texten."},
                    {"role": "user", "content": edit_prompt}
                ],
                temperature=0.5,
                max_tokens=2000,
                response_format={"type": "json_object"}
            )
            
            result = json.loads(response.choices[0].message.content)
            result["edit_success"] = True
            return result
            
        except Exception as e:
            logger.error(f"Edit error: {e}")
            return {
                "edit_success": False,
                "error": str(e)
            }
    
    @staticmethod
    async def suggest_improvements(
        question: Dict[str, Any],
        context: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Suggest improvements for a specific question"""
        
        prompt = f"""Analysiere diese Aufgabe und schlage Verbesserungen vor:

Aufgabe:
{json.dumps(question, ensure_ascii=False, indent=2)}

Kontext:
- Klasse: {context.get('grade', '4')}
- Fach: {context.get('subject', 'Deutsch')}
- Schwierigkeit: {context.get('difficulty', 'medium')}

Gib Vorschläge im JSON-Format:
{{
  "suggestions": [
    {{
      "type": "simplify|extend|clarify|alternative",
      "description": "Beschreibung der Verbesserung",
      "preview": "Vorschau des verbesserten Texts"
    }}
  ],
  "quality_score": 1-10,
  "issues": ["Erkannte Probleme"]
}}"""

        try:
            response = openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "Du bist ein Experte für Unterrichtsmaterialien."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.5,
                max_tokens=1000,
                response_format={"type": "json_object"}
            )
            
            return json.loads(response.choices[0].message.content)
            
        except Exception as e:
            logger.error(f"Suggestion error: {e}")
            return {"suggestions": [], "error": str(e)}
