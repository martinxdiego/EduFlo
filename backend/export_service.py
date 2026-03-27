"""
EduFlow Export Service
Handles professional PDF and Word document generation
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm, mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from io import BytesIO
import logging

logger = logging.getLogger(__name__)

class WorksheetExporter:
    """
    Professional export service for worksheets and exams
    """
    
    def __init__(self, worksheet_data, include_solutions=False):
        self.worksheet = worksheet_data
        self.include_solutions = include_solutions
        self.mode = worksheet_data.get('mode', 'worksheet')
        self.content = worksheet_data.get('content', {})
        
    def export_to_docx(self) -> BytesIO:
        """
        Export to Word (.docx) format
        Professional, editable template with answer lines
        """
        doc = Document()
        
        # Set page margins (A4)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Title
        title = doc.add_heading(self.worksheet.get('title', 'Arbeitsblatt'), level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add student name/date line for student version
        if not self.include_solutions:
            name_para = doc.add_paragraph()
            name_para.add_run('Name: ________________________    Datum: ____________')
            name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Metadata
        metadata = doc.add_paragraph()
        metadata.add_run(f"Klasse: {self.worksheet.get('grade')}   |   ")
        metadata.add_run(f"Fach: {self.worksheet.get('subject')}   |   ")
        difficulty_map = {'easy': 'Einfach', 'medium': 'Mittel', 'hard': 'Schwierig'}
        metadata.add_run(f"Schwierigkeit: {difficulty_map.get(self.worksheet.get('difficulty'), self.worksheet.get('difficulty'))}")
        metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Total points for exam mode (at top)
        if self.mode == 'exam':
            total_points = self.content.get('total_points', 0)
            if not total_points:
                # Calculate if not present
                total_points = sum(q.get('points', 0) for q in self.content.get('questions', []))
            
            points_para = doc.add_paragraph()
            points_run = points_para.add_run(f"Gesamtpunkte: _____ / {total_points}")
            points_run.bold = True
            points_run.font.size = Pt(14)
            points_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()  # Spacer
        
        # Questions
        for q in self.content.get('questions', []):
            # Question number and text
            q_para = doc.add_paragraph()
            q_num = q_para.add_run(f"{q.get('number', 1)}. ")
            q_num.bold = True
            q_num.font.size = Pt(11)
            
            q_text = q_para.add_run(q.get('question', ''))
            q_text.font.size = Pt(11)
            
            # Points (exam mode only, hide in worksheet mode)
            if self.mode == 'exam' and q.get('points'):
                points_run = q_para.add_run(f"  ({q.get('points')} Punkt{'e' if q.get('points', 0) > 1 else ''})")
                points_run.font.italic = True
                points_run.font.color.rgb = RGBColor(100, 100, 100)
            
            # Options (multiple choice)
            q_type = q.get('type', 'short_answer')
            if q.get('options') and q_type == 'multiple_choice':
                for option in q['options']:
                    option_para = doc.add_paragraph()
                    option_para.add_run(f"    ○  {option}")
                    option_para.paragraph_format.space_after = Pt(4)
            
            # Answer space (if not showing solutions)
            if not self.include_solutions:
                answer_lines = q.get('answerLines', 3)
                if q_type != 'multiple_choice' and answer_lines > 0:
                    # Add answer lines
                    for _ in range(answer_lines):
                        line_para = doc.add_paragraph()
                        line_para.add_run('_' * 70)
                        line_para.paragraph_format.space_before = Pt(8)
                        line_para.paragraph_format.space_after = Pt(2)
            
            # Solution (if requested - teacher version)
            if self.include_solutions and q.get('answer'):
                answer_para = doc.add_paragraph()
                answer_label = answer_para.add_run('Lösung: ')
                answer_label.bold = True
                answer_label.font.color.rgb = RGBColor(0, 128, 0)
                answer_text = answer_para.add_run(q['answer'])
                answer_text.font.color.rgb = RGBColor(0, 128, 0)
                answer_para.paragraph_format.left_indent = Inches(0.25)
                
                # Add explanation if available
                if q.get('explanation'):
                    exp_para = doc.add_paragraph()
                    exp_label = exp_para.add_run('Erklärung: ')
                    exp_label.italic = True
                    exp_para.add_run(q['explanation'])
                    exp_para.paragraph_format.left_indent = Inches(0.25)
            
            doc.add_paragraph()  # Spacer between questions
        
        # Teacher notes (only in teacher version)
        if self.include_solutions and self.content.get('teacher_notes'):
            doc.add_page_break()
            doc.add_heading('Lehrernotizen', level=2)
            notes_para = doc.add_paragraph(self.content['teacher_notes'])
            notes_para.paragraph_format.line_spacing = 1.5
        
        # Estimated time
        if self.content.get('estimated_time'):
            time_para = doc.add_paragraph()
            time_para.add_run(f"Geschätzte Zeit: {self.content['estimated_time']}")
            time_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        logger.info(f"Generated DOCX export: {self.worksheet.get('id')} (solutions: {self.include_solutions})")
        return buffer
    
    def export_to_pdf(self) -> BytesIO:
        """
        Export to PDF format
        Professional A4 layout with proper typography and answer lines
        """
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            topMargin=1.5*cm,
            bottomMargin=1.5*cm,
            leftMargin=2*cm,
            rightMargin=2*cm
        )
        
        # Styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#1e40af'),
            spaceAfter=8,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        metadata_style = ParagraphStyle(
            'Metadata',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.grey,
            spaceAfter=15,
            alignment=TA_CENTER
        )
        
        name_style = ParagraphStyle(
            'NameLine',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=10,
            alignment=TA_LEFT
        )
        
        question_style = ParagraphStyle(
            'Question',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            leftIndent=0,
            fontName='Helvetica-Bold'
        )
        
        answer_style = ParagraphStyle(
            'Answer',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#059669'),
            leftIndent=15,
            spaceAfter=8
        )
        
        option_style = ParagraphStyle(
            'Option',
            parent=styles['Normal'],
            fontSize=10,
            leftIndent=25,
            spaceAfter=4
        )
        
        line_style = ParagraphStyle(
            'AnswerLine',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.grey,
            leftIndent=15,
            spaceBefore=6,
            spaceAfter=2
        )
        
        # Build document
        story = []
        
        # Title
        story.append(Paragraph(self.worksheet.get('title', 'Arbeitsblatt'), title_style))
        
        # Name/Date line for student version
        if not self.include_solutions:
            story.append(Paragraph('Name: ________________________    Datum: ____________', name_style))
        
        # Metadata
        difficulty_map = {'easy': 'Einfach', 'medium': 'Mittel', 'hard': 'Schwierig'}
        metadata_text = f"Klasse: {self.worksheet.get('grade')} | Fach: {self.worksheet.get('subject')} | Schwierigkeit: {difficulty_map.get(self.worksheet.get('difficulty'), self.worksheet.get('difficulty'))}"
        story.append(Paragraph(metadata_text, metadata_style))
        
        # Total points box (exam mode - at top)
        if self.mode == 'exam':
            total_points = self.content.get('total_points', 0)
            if not total_points:
                total_points = sum(q.get('points', 0) for q in self.content.get('questions', []))
            
            points_data = [[f"Punkte: _____ / {total_points}"]]
            points_table = Table(points_data, colWidths=[15*cm])
            points_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#dbeafe')),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#1e40af')),
                ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 12),
                ('BOX', (0, 0), (-1, -1), 2, colors.HexColor('#3b82f6')),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ]))
            story.append(points_table)
            story.append(Spacer(1, 0.5*cm))
        
        # Questions
        for i, q in enumerate(self.content.get('questions', []), 1):
            q_type = q.get('type', 'short_answer')
            
            # Question text
            question_text = f"<b>{q.get('number', i)}.</b> {q.get('question', '')}"
            
            # Add points for exam mode (hide in worksheet mode)
            if self.mode == 'exam' and q.get('points'):
                question_text += f" <i><font color='#6b7280'>({q.get('points')} Punkt{'e' if q.get('points', 0) > 1 else ''})</font></i>"
            
            story.append(Paragraph(question_text, question_style))
            
            # Options (multiple choice)
            if q.get('options') and q_type == 'multiple_choice':
                for option in q['options']:
                    story.append(Paragraph(f"○  {option}", option_style))
            
            # Answer space or solution
            if self.include_solutions and q.get('answer'):
                story.append(Paragraph(f"<b>Lösung:</b> {q['answer']}", answer_style))
                if q.get('explanation'):
                    story.append(Paragraph(f"<i>Erklärung: {q['explanation']}</i>", answer_style))
            else:
                # Add answer lines for student version (except for multiple choice)
                if q_type != 'multiple_choice':
                    answer_lines = q.get('answerLines', 3)
                    for _ in range(answer_lines):
                        story.append(Paragraph('_' * 80, line_style))
            
            story.append(Spacer(1, 0.4*cm))
        
        # Estimated time
        if self.content.get('estimated_time'):
            story.append(Spacer(1, 0.5*cm))
            time_style = ParagraphStyle(
                'TimeStyle',
                parent=styles['Normal'],
                fontSize=9,
                textColor=colors.grey,
                alignment=TA_RIGHT
            )
            story.append(Paragraph(f"Geschätzte Zeit: {self.content['estimated_time']}", time_style))
        
        # Teacher notes (teacher version only)
        if self.include_solutions and self.content.get('teacher_notes'):
            story.append(PageBreak())
            story.append(Paragraph('<b>Lehrernotizen</b>', title_style))
            story.append(Spacer(1, 0.3*cm))
            story.append(Paragraph(self.content['teacher_notes'], styles['Normal']))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        
        logger.info(f"Generated PDF export: {self.worksheet.get('id')} (solutions: {self.include_solutions})")
        return buffer


def export_worksheet_pdf(worksheet_data, include_solutions=False) -> BytesIO:
    """Convenience function for PDF export"""
    exporter = WorksheetExporter(worksheet_data, include_solutions)
    return exporter.export_to_pdf()


def export_worksheet_docx(worksheet_data, include_solutions=False) -> BytesIO:
    """Convenience function for DOCX export"""
    exporter = WorksheetExporter(worksheet_data, include_solutions)
    return exporter.export_to_docx()


# ==================== DOSSIER EXPORTER ====================

# Theme color defaults (subset - frontend themes are defined in JS)
THEME_COLORS = {
    'classic': {'primary': '#1e40af', 'accent': '#3b82f6', 'light': '#dbeafe', 'line': '#d1d5db'},
    'abenteuer': {'primary': '#92400e', 'accent': '#d97706', 'light': '#fef3c7', 'line': '#d6d3d1'},
    'weltraum': {'primary': '#312e81', 'accent': '#6366f1', 'light': '#e0e7ff', 'line': '#c7d2fe'},
    'natur': {'primary': '#166534', 'accent': '#22c55e', 'light': '#dcfce7', 'line': '#bbf7d0'},
    'ozean': {'primary': '#0c4a6e', 'accent': '#0284c7', 'light': '#e0f2fe', 'line': '#bae6fd'},
    'bunt': {'primary': '#be185d', 'accent': '#ec4899', 'light': '#fce7f3', 'line': '#f9a8d4'},
    'comic': {'primary': '#dc2626', 'accent': '#f59e0b', 'light': '#fee2e2', 'line': '#fde68a'},
    'jahreszeit_fruehling': {'primary': '#db2777', 'accent': '#f472b6', 'light': '#fce7f3', 'line': '#fbcfe8'},
    'jahreszeit_herbst': {'primary': '#9a3412', 'accent': '#ea580c', 'light': '#ffedd5', 'line': '#fed7aa'},
    'jahreszeit_winter': {'primary': '#1e3a5f', 'accent': '#38bdf8', 'light': '#e0f2fe', 'line': '#bae6fd'},
}

INFO_BOX_STYLES = {
    'wusstest_du': {'bg': '#eff6ff', 'border': '#3b82f6', 'icon': 'Wusstest du?', 'text_color': '#1e40af'},
    'wichtig': {'bg': '#fef2f2', 'border': '#ef4444', 'icon': 'Wichtig!', 'text_color': '#991b1b'},
    'merke': {'bg': '#f0fdf4', 'border': '#22c55e', 'icon': 'Merke dir', 'text_color': '#166534'},
    'tipp': {'bg': '#fefce8', 'border': '#eab308', 'icon': 'Tipp', 'text_color': '#854d0e'},
}


class DossierExporter:
    """
    Professional PDF export for multi-page learning dossiers
    with title page, TOC, running headers, and themed blocks
    """

    def __init__(self, dossier_data, include_solutions=False):
        self.dossier = dossier_data
        self.include_solutions = include_solutions
        self.sections = dossier_data.get('sections', [])
        theme_id = dossier_data.get('theme', 'classic')
        self.theme = THEME_COLORS.get(theme_id, THEME_COLORS['classic'])
        self.page_width = A4[0] - 4 * cm  # usable width with margins
        self._setup_styles()

    def _setup_styles(self):
        """Setup all paragraph styles"""
        base = getSampleStyleSheet()
        tc = self.theme

        self.styles = {
            'title': ParagraphStyle('DTitle', parent=base['Heading1'], fontSize=24,
                textColor=colors.HexColor(tc['primary']), spaceAfter=10,
                alignment=TA_CENTER, fontName='Helvetica-Bold'),
            'subtitle': ParagraphStyle('DSubtitle', parent=base['Normal'], fontSize=12,
                textColor=colors.grey, spaceAfter=20, alignment=TA_CENTER),
            'h1': ParagraphStyle('DH1', parent=base['Heading1'], fontSize=18,
                textColor=colors.HexColor(tc['primary']), spaceBefore=20, spaceAfter=10,
                fontName='Helvetica-Bold'),
            'h2': ParagraphStyle('DH2', parent=base['Heading2'], fontSize=14,
                textColor=colors.HexColor(tc['primary']), spaceBefore=14, spaceAfter=8,
                fontName='Helvetica-Bold'),
            'h3': ParagraphStyle('DH3', parent=base['Heading3'], fontSize=12,
                textColor=colors.HexColor(tc['accent']), spaceBefore=10, spaceAfter=6,
                fontName='Helvetica-Bold'),
            'body': ParagraphStyle('DBody', parent=base['Normal'], fontSize=10,
                spaceAfter=6, leading=14),
            'question': ParagraphStyle('DQuestion', parent=base['Normal'], fontSize=10,
                spaceAfter=6, fontName='Helvetica-Bold'),
            'option': ParagraphStyle('DOption', parent=base['Normal'], fontSize=10,
                leftIndent=20, spaceAfter=3),
            'answer_line': ParagraphStyle('DLine', parent=base['Normal'], fontSize=10,
                textColor=colors.HexColor('#999999'), leftIndent=10, spaceBefore=5, spaceAfter=2),
            'solution': ParagraphStyle('DSolution', parent=base['Normal'], fontSize=10,
                textColor=colors.HexColor('#059669'), leftIndent=10, spaceAfter=6),
            'toc_entry': ParagraphStyle('DToc', parent=base['Normal'], fontSize=11,
                spaceAfter=6, leftIndent=10),
            'footer': ParagraphStyle('DFooter', parent=base['Normal'], fontSize=8,
                textColor=colors.grey, alignment=TA_CENTER),
            'objective': ParagraphStyle('DObjective', parent=base['Normal'], fontSize=10,
                leftIndent=20, spaceAfter=4),
            'glossary_term': ParagraphStyle('DGlossaryTerm', parent=base['Normal'], fontSize=10,
                fontName='Helvetica-Bold', spaceAfter=2),
            'glossary_def': ParagraphStyle('DGlossaryDef', parent=base['Normal'], fontSize=10,
                leftIndent=15, spaceAfter=8),
            'reflection': ParagraphStyle('DReflection', parent=base['Normal'], fontSize=10,
                leftIndent=15, spaceBefore=4, spaceAfter=4,
                textColor=colors.HexColor(tc['primary'])),
        }

    def _header_footer(self, canvas, doc):
        """Add running header and footer to each page"""
        canvas.saveState()
        tc = self.theme

        # Header line
        canvas.setStrokeColor(colors.HexColor(tc['line']))
        canvas.setLineWidth(0.5)
        canvas.line(2*cm, A4[1] - 1.2*cm, A4[0] - 2*cm, A4[1] - 1.2*cm)

        # Header text
        canvas.setFont('Helvetica', 8)
        canvas.setFillColor(colors.HexColor('#999999'))
        header_left = f"{self.dossier.get('subject', '')} | {self.dossier.get('topic', '')}"
        canvas.drawString(2*cm, A4[1] - 1.1*cm, header_left)

        # Page number
        page_num = f"Seite {doc.page}"
        canvas.drawRightString(A4[0] - 2*cm, A4[1] - 1.1*cm, page_num)

        # Footer line
        canvas.line(2*cm, 1.2*cm, A4[0] - 2*cm, 1.2*cm)
        canvas.drawCentredString(A4[0] / 2, 0.8*cm, self.dossier.get('title', ''))

        canvas.restoreState()

    def _first_page(self, canvas, doc):
        """Title page - no header/footer"""
        pass  # Title page has no running header

    def _render_block(self, block, story, question_counter):
        """Render a single block into story elements"""
        block_type = block.get('type', 'text')
        content = block.get('content', {})

        if block_type == 'heading':
            level = content.get('level', 2)
            style_key = f'h{min(level, 3)}'
            text = content.get('text', '')
            story.append(Paragraph(text, self.styles[style_key]))

        elif block_type == 'text':
            html = content.get('html', '')
            if html:
                story.append(Paragraph(html, self.styles['body']))

        elif block_type == 'info_box':
            variant = content.get('variant', 'tipp')
            box_style = INFO_BOX_STYLES.get(variant, INFO_BOX_STYLES['tipp'])
            title = content.get('title', box_style['icon'])
            box_content = content.get('content', '')

            # Build info box as a colored table
            box_text = f"<b>{title}</b><br/>{box_content}"
            box_para = Paragraph(box_text, ParagraphStyle('InfoBoxText',
                parent=self.styles['body'],
                textColor=colors.HexColor(box_style['text_color'])))

            box_table = Table([[box_para]], colWidths=[self.page_width - 0.5*cm])
            box_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(box_style['bg'])),
                ('BOX', (0, 0), (-1, -1), 2, colors.HexColor(box_style['border'])),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('LEFTPADDING', (0, 0), (-1, -1), 12),
                ('RIGHTPADDING', (0, 0), (-1, -1), 12),
            ]))
            story.append(box_table)
            story.append(Spacer(1, 0.3*cm))

        elif block_type == 'question':
            question_counter[0] += 1
            q_num = content.get('number', question_counter[0])
            q_text = content.get('question', '')
            q_type = content.get('type', 'open')

            story.append(Paragraph(f"<b>{q_num}.</b> {q_text}", self.styles['question']))

            # Multiple choice options
            if q_type == 'multiple_choice' and content.get('options'):
                for opt in content['options']:
                    story.append(Paragraph(f"○  {opt}", self.styles['option']))

            # Answer space or solution
            if self.include_solutions and content.get('answer'):
                story.append(Paragraph(f"<b>Loesung:</b> {content['answer']}", self.styles['solution']))
                if content.get('explanation'):
                    story.append(Paragraph(f"<i>{content['explanation']}</i>", self.styles['solution']))
            elif q_type != 'multiple_choice':
                answer_lines = content.get('answerLines', 3)
                for _ in range(answer_lines):
                    story.append(Paragraph('_' * 80, self.styles['answer_line']))

            story.append(Spacer(1, 0.3*cm))

        elif block_type == 'table':
            headers = content.get('headers', [])
            rows = content.get('rows', [])
            if headers or rows:
                table_data = []
                if headers:
                    table_data.append(headers)
                table_data.extend(rows)

                col_count = max(len(r) for r in table_data) if table_data else 1
                col_width = self.page_width / col_count

                t = Table(table_data, colWidths=[col_width] * col_count)
                style_cmds = [
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor(self.theme['line'])),
                    ('TOPPADDING', (0, 0), (-1, -1), 6),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ('LEFTPADDING', (0, 0), (-1, -1), 8),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                ]
                if headers:
                    style_cmds.extend([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(self.theme['light'])),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor(self.theme['primary'])),
                    ])
                t.setStyle(TableStyle(style_cmds))
                story.append(t)
                story.append(Spacer(1, 0.3*cm))

        elif block_type == 'objectives_checklist':
            objectives = content.get('objectives', [])
            for obj in objectives:
                code = obj.get('code', '')
                text = obj.get('text', '')
                prefix = f"<b>{code}</b> " if code else ""
                story.append(Paragraph(f"☐ {prefix}{text}", self.styles['objective']))

        elif block_type == 'glossary':
            terms = content.get('terms', [])
            for item in terms:
                term = item.get('term', '')
                definition = item.get('definition', '')
                story.append(Paragraph(term, self.styles['glossary_term']))
                story.append(Paragraph(definition, self.styles['glossary_def']))

        elif block_type == 'reflection':
            prompts = content.get('prompts', [])
            for prompt in prompts:
                story.append(Paragraph(f"• {prompt}", self.styles['reflection']))
                # Add a few lines for writing
                for _ in range(2):
                    story.append(Paragraph('_' * 80, self.styles['answer_line']))
                story.append(Spacer(1, 0.2*cm))

        elif block_type == 'creative_task':
            instruction = content.get('instruction', '')
            space_lines = content.get('space_lines', 8)
            story.append(Paragraph(instruction, self.styles['body']))
            story.append(Spacer(1, 0.2*cm))
            for _ in range(space_lines):
                story.append(Paragraph('_' * 80, self.styles['answer_line']))

        elif block_type == 'image_placeholder':
            desc = content.get('description', 'Bild')
            caption = content.get('caption', '')
            placeholder_text = f"[Bild: {desc}]"
            if caption:
                placeholder_text += f"<br/><i>{caption}</i>"

            box = Table([[Paragraph(placeholder_text, ParagraphStyle('ImgPlaceholder',
                parent=self.styles['body'], alignment=TA_CENTER,
                textColor=colors.HexColor('#999999')))]],
                colWidths=[self.page_width * 0.7])
            box.setStyle(TableStyle([
                ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#cccccc')),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f9fafb')),
                ('TOPPADDING', (0, 0), (-1, -1), 20),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 20),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ]))
            story.append(box)
            story.append(Spacer(1, 0.3*cm))

        elif block_type == 'page_break':
            story.append(PageBreak())

        elif block_type == 'divider':
            story.append(Spacer(1, 0.3*cm))
            # Simple line divider
            divider = Table([['']],  colWidths=[self.page_width])
            divider.setStyle(TableStyle([
                ('LINEABOVE', (0, 0), (-1, 0), 0.5, colors.HexColor(self.theme['line'])),
            ]))
            story.append(divider)
            story.append(Spacer(1, 0.3*cm))

        return question_counter

    def export_to_pdf(self) -> BytesIO:
        """Export the complete dossier to PDF"""
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            topMargin=1.5*cm,
            bottomMargin=1.5*cm,
            leftMargin=2*cm,
            rightMargin=2*cm
        )

        story = []
        tc = self.theme
        difficulty_map = {'easy': 'Einfach', 'medium': 'Mittel', 'hard': 'Schwierig'}

        # ===== TITLE PAGE =====
        story.append(Spacer(1, 3*cm))

        # Accent bar
        bar = Table([['']],  colWidths=[self.page_width])
        bar.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(tc['accent'])),
            ('TOPPADDING', (0, 0), (-1, -1), 3),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ]))
        story.append(bar)
        story.append(Spacer(1, 1*cm))

        # Title
        story.append(Paragraph(self.dossier.get('title', 'Lerndossier'), self.styles['title']))
        story.append(Spacer(1, 0.5*cm))

        # Metadata
        grade = self.dossier.get('grade', '')
        subject = self.dossier.get('subject', '')
        difficulty = difficulty_map.get(self.dossier.get('difficulty', ''), '')
        meta_text = f"{subject} | {grade}. Klasse | {difficulty}"
        story.append(Paragraph(meta_text, self.styles['subtitle']))

        # Name/Date line
        story.append(Spacer(1, 2*cm))
        story.append(Paragraph('Name: ________________________    Datum: ____________',
            ParagraphStyle('TitleName', parent=self.styles['body'], fontSize=11)))

        story.append(Spacer(1, 1*cm))
        story.append(bar)  # Bottom accent bar

        story.append(PageBreak())

        # ===== TABLE OF CONTENTS =====
        story.append(Paragraph('Inhaltsverzeichnis', self.styles['h1']))
        story.append(Spacer(1, 0.5*cm))

        for idx, section in enumerate(self.sections):
            if section.get('type') == 'solutions' and not self.include_solutions:
                continue
            section_title = section.get('title', f'Sektion {idx + 1}')
            story.append(Paragraph(
                f"<b>{idx + 1}.</b>  {section_title}",
                self.styles['toc_entry']
            ))

        story.append(PageBreak())

        # ===== SECTIONS =====
        question_counter = [0]  # mutable counter for questions across sections

        for idx, section in enumerate(self.sections):
            section_type = section.get('type', 'theory')
            section_title = section.get('title', '')

            # Skip solutions section in student version
            if section_type == 'solutions' and not self.include_solutions:
                continue

            # Section heading
            story.append(Paragraph(f"{idx + 1}. {section_title}", self.styles['h1']))

            # Accent line under section title
            accent_line = Table([['']],  colWidths=[self.page_width])
            accent_line.setStyle(TableStyle([
                ('LINEBELOW', (0, 0), (-1, 0), 1.5, colors.HexColor(tc['accent'])),
            ]))
            story.append(accent_line)
            story.append(Spacer(1, 0.4*cm))

            # Render blocks
            for block in section.get('blocks', []):
                question_counter = self._render_block(block, story, question_counter)

            # Page break between sections (except last)
            if idx < len(self.sections) - 1:
                story.append(PageBreak())

        # Build PDF with header/footer
        doc.build(story, onFirstPage=self._first_page, onLaterPages=self._header_footer)
        buffer.seek(0)

        logger.info(f"Generated dossier PDF: {self.dossier.get('id')} ({len(self.sections)} sections, solutions: {self.include_solutions})")
        return buffer


def export_dossier_pdf(dossier_data, include_solutions=False) -> BytesIO:
    """Convenience function for dossier PDF export"""
    exporter = DossierExporter(dossier_data, include_solutions)
    return exporter.export_to_pdf()
